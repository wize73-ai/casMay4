"""
DOCX Processor for CasaLingua

This module extracts text content from DOCX files and provides basic preprocessing.

Author: Exygy Development Team
Version: 1.0.0
License: MIT
"""

import logging
from typing import Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.utils.logging import get_logger

logger = get_logger(__name__)

class DOCXProcessor:
    """
    DOCXProcessor handles the loading and extraction of text from DOCX documents.
    """

    def __init__(self, model_manager=None, config=None):
        self.model_manager = model_manager
        self.config = config

    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from a DOCX file.

        Args:
            file_path (str): Path to the DOCX file.

        Returns:
            str or None: Extracted text content, or None if loading failed.
        """
        try:
            document = Document(file_path)
            text = "\n".join([para.text for para in document.paragraphs])
            logger.debug(f"Extracted {len(text)} characters from DOCX file {file_path}")
            return text
        except PackageNotFoundError as e:
            logger.error(f"Invalid DOCX file at {file_path}: {e}")
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
        return None

    def create_document(self, text: str, output_path: str) -> bool:
        """
        Create a new DOCX document with the given text content.

        Args:
            text (str): Text to include in the document.
            output_path (str): Path to save the generated DOCX file.

        Returns:
            bool: True if document was successfully saved, False otherwise.
        """
        try:
            doc = Document()
            for line in text.strip().split("\n"):
                doc.add_paragraph(line.strip())
            doc.save(output_path)
            logger.info(f"DOCX file created at {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to create DOCX file: {e}", exc_info=True)
            return False

if __name__ == "__main__":
    import os

    logging.basicConfig(level=logging.DEBUG)
    processor = DOCXProcessor()

    test_text = "CasaLingua DOCX Test\nLine 2 of content."
    test_output_path = "test_output.docx"

    # Create document
    success = processor.create_document(test_text, test_output_path)
    if success and os.path.exists(test_output_path):
        print(f"✅ DOCX created at {test_output_path}")
    else:
        print("❌ DOCX creation failed.")

    # Extract document
    extracted = processor.extract_text(test_output_path)
    if extracted:
        print("✅ DOCX text extracted:")
        print(extracted)
    else:
        print("❌ Failed to extract DOCX text.")