"""
Indexer Module for CasaLingua
Provides functionality for indexing, chunking, and preparing documents for the RAG system
"""

import os
import json
import hashlib
import re
import logging
from typing import List, Dict, Union, Tuple, Optional, Any, Callable, Iterator
from datetime import datetime
from docx import Document
import pandas as pd
from tqdm import tqdm

# Tokenizer pipeline import
from app.core.pipeline.tokenizer import TokenizerPipeline
# ModelRegistry import for dynamic tokenizer loading
from app.services.models.loader import ModelRegistry

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Base class for document processors"""
    
    def __init__(self):
        """Initialize the document processor"""
        pass
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a document into chunks suitable for indexing
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        raise NotImplementedError("Subclasses must implement process()")
    
    @staticmethod
    def detect_language(text: str) -> str:
        """
        Simple keyword-based heuristic for language detection.
        Replace with langdetect or spaCy for production quality.

        Args:
            text (str): Text to detect language for

        Returns:
            str: Detected language code (e.g., 'en', 'es', 'fr')
        """
        text = text.lower()
        language_markers = {
            'en': ['the', 'and', 'of', 'to', 'in', 'is', 'you', 'that', 'it', 'he'],
            'es': ['el', 'la', 'de', 'que', 'y', 'a', 'en', 'un', 'ser', 'se'],
            'fr': ['le', 'la', 'de', 'et', 'est', 'en', 'que', 'un', 'une', 'du'],
            'de': ['der', 'die', 'das', 'und', 'ist', 'in', 'zu', 'den', 'mit', 'nicht'],
            'it': ['il', 'la', 'di', 'e', 'che', 'un', 'a', 'per', 'in', 'sono']
        }
        scores = {lang: 0 for lang in language_markers}
        for lang, markers in language_markers.items():
            for marker in markers:
                pattern = r'\b' + re.escape(marker) + r'\b'
                matches = re.findall(pattern, text)
                scores[lang] += len(matches)
        max_score = max(scores.values())
        if max_score == 0:
            return 'unknown'
        return max(scores.items(), key=lambda x: x[1])[0]


class TextProcessor(DocumentProcessor):
    """Processor for plain text files"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50, tokenizer: Optional[TokenizerPipeline] = None):
        """
        Initialize the text processor
        
        Args:
            chunk_size (int): Maximum size of text chunks in characters
            chunk_overlap (int): Overlap between consecutive chunks in characters
            tokenizer (TokenizerPipeline, optional): Shared tokenizer pipeline
        """
        super().__init__()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tokenizer
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a text file into chunks
        
        Args:
            file_path (str): Path to the text file
            
        Returns:
            List[Dict[str, Any]]: List of text chunks with metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            return self._chunk_text(text, os.path.basename(file_path))
                
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return []
    
    def _chunk_text(self, text: str, source: str) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks
        
        Args:
            text (str): Text to split
            source (str): Source identifier
            
        Returns:
            List[Dict[str, Any]]: List of text chunks with metadata
        """
        chunks = []
        
        # Simple chunking by character count
        # For better results, consider chunking by sentences or paragraphs
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # If we're not at the end, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence boundaries (., !, ?)
                last_period = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end)
                )
                
                if last_period != -1:
                    end = last_period + 1  # Include the period
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:  # Only add non-empty chunks
                # Detect language
                language = self.detect_language(chunk_text)
                
                # Create unique ID for the chunk
                chunk_id = hashlib.md5(chunk_text.encode('utf-8')).hexdigest()
                chunk = {
                    "id": chunk_id,
                    "text": chunk_text,
                    "source": source,
                    "metadata": {
                        "language": language,
                        "start_char": start,
                        "end_char": end,
                        "character_count": len(chunk_text)
                    }
                }
                if self.tokenizer:
                    chunk_tokens = self.tokenizer.encode(chunk_text)
                    chunk["metadata"]["tokens"] = chunk_tokens
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap if end < len(text) else end
        
        return chunks


class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents"""
    
    def __init__(self, max_paragraphs_per_chunk: int = 5, tokenizer: Optional[TokenizerPipeline] = None):
        """
        Initialize the Word document processor
        
        Args:
            max_paragraphs_per_chunk (int): Maximum number of paragraphs per chunk
            tokenizer (TokenizerPipeline, optional): Shared tokenizer pipeline
        """
        super().__init__()
        self.max_paragraphs_per_chunk = max_paragraphs_per_chunk
        self.tokenizer = tokenizer
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a Word document into chunks
        
        Args:
            file_path (str): Path to the Word document
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        try:
            doc = Document(file_path)
            
            chunks = []
            current_paragraphs = []
            current_text = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Skip empty paragraphs
                    current_paragraphs.append(paragraph.text)
                    current_text += paragraph.text + "\n"
                    
                    # Create chunk when we reach max paragraphs
                    if len(current_paragraphs) >= self.max_paragraphs_per_chunk:
                        chunk_text = "\n".join(current_paragraphs)
                        
                        # Detect language
                        language = self.detect_language(chunk_text)
                        
                        # Create unique ID for the chunk
                        chunk_id = hashlib.md5(chunk_text.encode('utf-8')).hexdigest()
                        chunk = {
                            "id": chunk_id,
                            "text": chunk_text,
                            "source": os.path.basename(file_path),
                            "metadata": {
                                "language": language,
                                "paragraph_index_range": f"{i-len(current_paragraphs)+1}-{i}",
                                "paragraph_count": len(current_paragraphs),
                                "character_count": len(chunk_text)
                            }
                        }
                        if self.tokenizer:
                            chunk_tokens = self.tokenizer.encode(chunk_text)
                            chunk["metadata"]["tokens"] = chunk_tokens
                        chunks.append(chunk)
                        
                        current_paragraphs = []
                        current_text = ""
            
            # Add remaining paragraphs as a chunk
            if current_paragraphs:
                chunk_text = "\n".join(current_paragraphs)
                
                # Detect language
                language = self.detect_language(chunk_text)
                
                # Create unique ID for the chunk
                chunk_id = hashlib.md5(chunk_text.encode('utf-8')).hexdigest()
                chunk = {
                    "id": chunk_id,
                    "text": chunk_text,
                    "source": os.path.basename(file_path),
                    "metadata": {
                        "language": language,
                        "paragraph_index_range": f"{len(doc.paragraphs)-len(current_paragraphs)+1}-{len(doc.paragraphs)}",
                        "paragraph_count": len(current_paragraphs),
                        "character_count": len(chunk_text)
                    }
                }
                if self.tokenizer:
                    chunk_tokens = self.tokenizer.encode(chunk_text)
                    chunk["metadata"]["tokens"] = chunk_tokens
                chunks.append(chunk)
            
            return chunks
                
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            return []


class CSVProcessor(DocumentProcessor):
    """Processor for CSV files"""
    
    def __init__(self, text_columns: Optional[List[str]] = None, tokenizer: Optional[TokenizerPipeline] = None):
        """
        Initialize the CSV processor
        
        Args:
            text_columns (List[str], optional): List of column names to include
            tokenizer (TokenizerPipeline, optional): Shared tokenizer pipeline
        """
        super().__init__()
        self.text_columns = text_columns
        self.tokenizer = tokenizer
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a CSV file into document chunks
        
        Args:
            file_path (str): Path to the CSV file
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # If no specific text columns provided, use all columns
            text_columns = self.text_columns or df.columns.tolist()
            
            chunks = []
            
            # Process each row
            for i, row in df.iterrows():
                # Build text from specified columns
                row_text = " ".join(str(row[col]) for col in text_columns if col in row)
                
                if row_text.strip():
                    # Detect language
                    language = self.detect_language(row_text)
                    
                    # Create unique ID for the chunk
                    chunk_id = hashlib.md5(row_text.encode('utf-8')).hexdigest()
                    
                    # Standardized metadata keys
                    metadata = {
                        "row_index": i,
                        "language": language,
                        "character_count": len(row_text)
                    }
                    # Add any additional columns as metadata
                    for col in df.columns:
                        if col not in text_columns and pd.notna(row[col]):
                            metadata[col] = row[col]
                    if self.tokenizer:
                        chunk_tokens = self.tokenizer.encode(row_text)
                        metadata["tokens"] = chunk_tokens
                    chunks.append({
                        "id": chunk_id,
                        "text": row_text,
                        "source": os.path.basename(file_path),
                        "metadata": metadata
                    })
            
            return chunks
                
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            return []


class Indexer:
    """Main indexer class for CasaLingua"""

    def __init__(self,
                output_dir: str = "./index",
                chunk_size: int = 500,
                chunk_overlap: int = 50):
        """
        Initialize the indexer

        Args:
            output_dir (str): Directory to save indexed documents
            chunk_size (int): Default chunk size for text processors
            chunk_overlap (int): Default chunk overlap for text processors
        """
        self.output_dir = output_dir
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Load tokenizer dynamically from registry for RAG retriever task
        registry = ModelRegistry()
        _, tokenizer_name = registry.get_model_and_tokenizer("rag_retriever")
        tokenizer = TokenizerPipeline(model_name=tokenizer_name, task_type="rag_retrieval")

        # Initialize document processors with dynamically loaded tokenizer
        self.processors = {
            ".txt": TextProcessor(chunk_size, chunk_overlap, tokenizer=tokenizer),
            ".docx": DocxProcessor(tokenizer=tokenizer),
            ".csv": CSVProcessor(tokenizer=tokenizer)
        }

        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
    
    def register_processor(self, extension: str, processor: DocumentProcessor) -> None:
        """
        Register a custom document processor
        
        Args:
            extension (str): File extension to associate with processor
            processor (DocumentProcessor): Processor instance
        """
        self.processors[extension.lower()] = processor
    
    def index_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Index a single file
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            List[Dict[str, Any]]: List of indexed document chunks
        """
        if not os.path.isfile(file_path):
            logger.warning(f"File does not exist: {file_path}")
            return []
            
        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()
        
        # Check if we have a processor for this extension
        if ext not in self.processors:
            logger.warning(f"No processor registered for extension {ext}")
            return []
            
        # Process the file
        logger.info(f"Indexing file: {file_path}")
        processor = self.processors[ext]
        chunks = processor.process(file_path)
        
        # Add indexing metadata
        for chunk in chunks:
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            chunk["metadata"].update({
                "indexed_at": datetime.now().isoformat(),
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_extension": ext
            })
        
        return chunks
    
    def index_directory(self, 
                      directory: str, 
                      recursive: bool = True,
                      file_extensions: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Index all files in a directory
        
        Args:
            directory (str): Directory path
            recursive (bool): Whether to process subdirectories
            file_extensions (List[str], optional): Only process these extensions
            
        Returns:
            List[Dict[str, Any]]: List of all indexed document chunks
        """
        if not os.path.isdir(directory):
            logger.warning(f"Directory does not exist: {directory}")
            return []
            
        all_chunks = []
        
        # Get list of files
        files = []
        if recursive:
            for root, _, filenames in os.walk(directory):
                for filename in filenames:
                    files.append(os.path.join(root, filename))
        else:
            files = [os.path.join(directory, f) for f in os.listdir(directory) 
                    if os.path.isfile(os.path.join(directory, f))]
        
        # Filter by extension if needed
        if file_extensions:
            filtered_extensions = [ext.lower() if ext.startswith('.') else f".{ext.lower()}" 
                                for ext in file_extensions]
            files = [f for f in files if os.path.splitext(f)[1].lower() in filtered_extensions]
        
        # Process each file
        for file_path in tqdm(files, desc="📁 Indexing files"):
            chunks = self.index_file(file_path)
            all_chunks.extend(chunks)
        
        return all_chunks
    
    def save_index(self, chunks: List[Dict[str, Any]], output_file: Optional[str] = None) -> str:
        """
        Save indexed chunks to a file
        
        Args:
            chunks (List[Dict[str, Any]]): List of document chunks
            output_file (str, optional): Output file path
            
        Returns:
            str: Path to the saved file
        """
        if not chunks:
            logger.warning("No chunks to save")
            return ""
            
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = os.path.join(self.output_dir, f"index_{timestamp}.json")
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(chunks, f, ensure_ascii=False, indent=2)
            logger.info(f"[Indexer] ✅ Saved {len(chunks)} chunks to '{output_file}'")
            return output_file
        except Exception as e:
            logger.error(f"Error saving index: {e}")
            return ""
    
    def index_and_save(self, 
                     path: str, 
                     output_file: Optional[str] = None,
                     recursive: bool = True,
                     file_extensions: Optional[List[str]] = None) -> Tuple[str, int]:
        """
        Index a file or directory and save the results
        
        Args:
            path (str): Path to file or directory
            output_file (str, optional): Output file path
            recursive (bool): Whether to process subdirectories
            file_extensions (List[str], optional): Only process these extensions
            
        Returns:
            Tuple[str, int]: Output file path and number of chunks indexed
        """
        chunks = []
        
        if os.path.isfile(path):
            chunks = self.index_file(path)
        elif os.path.isdir(path):
            chunks = self.index_directory(path, recursive, file_extensions)
        else:
            logger.warning(f"Path does not exist: {path}")
            return "", 0
        
        output_path = self.save_index(chunks, output_file)
        return output_path, len(chunks)


# Example usage
if __name__ == "__main__":
    # Create indexer with default settings
    indexer = Indexer(output_dir="./indexed_data")
    
    # Example 1: Index a single text file
    #file_chunks = indexer.index_file("path/to/language_lesson.txt")
    #indexer.save_index(file_chunks, "language_lesson_index.json")
    
    # Example 2: Index a directory
    #output_file, chunk_count = indexer.index_and_save(
    #    "path/to/language_materials",
    #    recursive=True,
    #    file_extensions=[".txt", ".docx", ".csv"]
    #)
    #print(f"Indexed {chunk_count} chunks, saved to {output_file}")
    
    # Example 3: Custom processor
    # class PDFProcessor(DocumentProcessor):
    #     def process(self, file_path: str) -> List[Dict[str, Any]]:
    #         # PDF processing logic here
    #         pass
    # 
    # indexer.register_processor(".pdf", PDFProcessor())
    
    print("Indexer module initialized. Use indexer.index_and_save() to begin indexing.")